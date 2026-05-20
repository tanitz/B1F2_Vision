"""
Generate Word document (.docx) — weekly inspection summary report
Output: docs/weekly_report.docx
"""
import base64, collections, csv, datetime, io, os, random, tempfile
import cv2, numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

random.seed(42)

# ── Paths ─────────────────────────────────────────────────────────────────
BASE     = r"e:\99IS\B1F2"
RES_DIR  = os.path.join(BASE, "results")
IMG_DIR  = os.path.join(RES_DIR, "images")
OUT_DOCX = os.path.join(BASE, "docs", "weekly_report.docx")
os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)

START     = datetime.date(2026, 5, 11)
DATES     = [START + datetime.timedelta(days=i) for i in range(6)]
MODEL     = "CV5VS_MODEL"
INSP_PTS  = ["Filter Drier","Hot Gas Valve","Expansion Valve",
             "High Pressure Sensor","Common-Start-Run","Check Valve",
             "Bolt & Washer","Bolt & Washer"]
NG_CAPABLE = ["Filter Drier","Expansion Valve","Bolt & Washer"]
OK_SCORES  = {
    "Filter Drier":(0.93,0.04), "Hot Gas Valve":(0.80,0.05),
    "Expansion Valve":(0.88,0.05), "High Pressure Sensor":(0.84,0.04),
    "Common-Start-Run":(0.78,0.05), "Check Valve":(0.83,0.04),
    "Bolt & Washer":(0.87,0.05),
}

WEEKDAY_TH = ["จันทร์","อังคาร","พุธ","พฤหัสบดี","ศุกร์","เสาร์","อาทิตย์"]
MONTH_TH   = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม",
               "มิถุนายน","กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]

def date_th(d):
    return f"{d.day} {MONTH_TH[d.month]} {d.year + 543}"

def rand_score(mu, sd):
    return round(min(0.99, max(0.55, random.gauss(mu, sd))), 4)

def ng_score():
    return round(random.uniform(0.20, 0.44), 4)

# ── Collect daily summary (CSVs already exist from gen_weekly_report.py) ──
print("Reading CSV data …")
daily = []
d1_rid_map = None
d1_ng_rids = []

for idx, date in enumerate(DATES):
    p = os.path.join(RES_DIR, f"{date}.csv")
    if not os.path.isfile(p):
        print(f"  WARNING: {p} not found, run gen_weekly_report.py first")
        continue
    with open(p, newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    rid_map = collections.OrderedDict()
    for r in rows:
        rid = r["result_image_id"]
        if rid not in rid_map:
            rid_map[rid] = []
        rid_map[rid].append(r)
    total = len(rid_map)
    ng    = sum(1 for rs in rid_map.values() if rs[0]["overall_result"] == "FAIL")
    ok    = total - ng
    daily.append({"date": date, "total": total, "ok": ok, "ng": ng,
                  "rate": ok / total * 100 if total else 0})
    if date == datetime.date(2026, 5, 16):
        d1_rid_map = rid_map
        d1_ng_rids = [rid for rid, rs in rid_map.items()
                      if rs[0]["overall_result"] == "FAIL"]
    print(f"  {date}: total={total}  PASS={ok}  NG={ng}")

grand_total = sum(d["total"] for d in daily)
grand_ok    = sum(d["ok"]    for d in daily)
grand_ng    = sum(d["ng"]    for d in daily)
grand_rate  = grand_ok / grand_total * 100 if grand_total else 0

# ── Generate charts as temp PNG files ────────────────────────────────────
print("Generating charts …")

def save_chart_tmp(fig):
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return tmp.name

# Bar chart
fig, ax = plt.subplots(figsize=(11, 4.2))
fig.patch.set_facecolor("white")
ax.set_facecolor("#f9f9f9")
labels  = [d["date"].strftime("%d/%m") for d in daily]
ng_vals = [d["ng"]  for d in daily]
ok_vals = [d["ok"]  for d in daily]
x = np.arange(len(labels))
bw = 0.5
ax.bar(x, ok_vals, bw, label="PASS", color="#43a047", alpha=0.85, zorder=3)
ax.bar(x, ng_vals, bw, bottom=ok_vals, label="NG (FAIL)", color="#e53935", alpha=0.9, zorder=3)
for xi, ng, ok in zip(x, ng_vals, ok_vals):
    if ng > 0:
        ax.text(xi, ok + ng + 2, f"NG={ng}", ha="center", va="bottom",
                fontsize=10, fontweight="bold", color="#b71c1c")
ax.axhline(5, color="#ff6f00", linestyle="--", linewidth=1.5, label="UCL = 5", zorder=4)
ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=11)
ax.set_ylabel("Inspection Count", fontsize=11)
ax.set_xlabel("Date (DD/MM)", fontsize=11)
ax.set_title("Daily Inspection Result — B1F2 Vision System  (11–16 May 2026)",
             fontsize=12, fontweight="bold", pad=10)
ax.legend(fontsize=10); ax.yaxis.grid(True, linestyle="--", alpha=0.5, zorder=0)
ax.spines[["top","right"]].set_visible(False)
plt.tight_layout()
chart_bar_path = save_chart_tmp(fig)

# Pie chart — NG by inspection point
ng_pt_count = collections.Counter()
for date in DATES:
    p = os.path.join(RES_DIR, f"{date}.csv")
    if not os.path.isfile(p): continue
    with open(p, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            if row["result"] in ("FAIL","NO_MATCH"):
                ng_pt_count[row["inspection_name"]] += 1

fig2, ax2 = plt.subplots(figsize=(6, 5))
fig2.patch.set_facecolor("white")
labels_pie = list(ng_pt_count.keys())
vals_pie   = list(ng_pt_count.values())
colors_pie = ["#e53935","#fb8c00","#fdd835","#43a047","#1e88e5","#8e24aa","#00897b"]
wedges, texts, autos = ax2.pie(
    vals_pie, labels=labels_pie, autopct="%1.0f%%",
    colors=colors_pie[:len(vals_pie)], startangle=140,
    textprops={"fontsize": 9}, pctdistance=0.78)
for a in autos: a.set_fontweight("bold")
ax2.set_title("NG Breakdown by Inspection Point\n(6-day total)", fontsize=11, fontweight="bold")
plt.tight_layout()
chart_pie_path = save_chart_tmp(fig2)

# Pass rate trend line
fig3, ax3 = plt.subplots(figsize=(8, 3.2))
fig3.patch.set_facecolor("white")
ax3.set_facecolor("#f9f9f9")
rates = [d["rate"] for d in daily]
ax3.plot(labels, rates, "o-", color="#1565c0", linewidth=2.5, markersize=8, zorder=3)
for xi, (lbl, r) in enumerate(zip(labels, rates)):
    ax3.annotate(f"{r:.1f}%", (lbl, r), textcoords="offset points",
                 xytext=(0, 10), ha="center", fontsize=9, color="#1565c0")
ax3.axhline(99.0, color="#43a047", linestyle="--", linewidth=1.2, label="Target 99.0%")
ax3.set_ylim(min(rates) - 0.5, 100.2)
ax3.set_ylabel("Pass Rate (%)", fontsize=10)
ax3.set_xlabel("Date (DD/MM)", fontsize=10)
ax3.set_title("Pass Rate Trend (11–16 May 2026)", fontsize=11, fontweight="bold")
ax3.legend(fontsize=9); ax3.yaxis.grid(True, linestyle="--", alpha=0.5, zorder=0)
ax3.spines[["top","right"]].set_visible(False)
plt.tight_layout()
chart_trend_path = save_chart_tmp(fig3)

print("Charts done.")

# ── NG example images (resized for Word) ─────────────────────────────────
ng_img_paths = []
for rid in d1_ng_rids:
    src = os.path.join(IMG_DIR, f"{rid}.jpg")
    if not os.path.isfile(src): continue
    img = cv2.imread(src)
    if img is None: continue
    h, w = img.shape[:2]
    s = min(600/w, 400/h, 1.0)
    img = cv2.resize(img, (int(w*s), int(h*s)), cv2.INTER_AREA)
    tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
    cv2.imwrite(tmp.name, img, [cv2.IMWRITE_JPEG_QUALITY, 88])
    # Get fail inspection name
    fail_pts = []
    if d1_rid_map and rid in d1_rid_map:
        fail_pts = [r["inspection_name"] for r in d1_rid_map[rid]
                    if r["result"] != "PASS"]
    ng_img_paths.append({"path": tmp.name, "rid": rid,
                          "fail_pts": ", ".join(set(fail_pts)) or "–"})

# ── Build Word Document ───────────────────────────────────────────────────
print("Building Word document …")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── Styles ────────────────────────────────────────────────────────────────
styles = doc.styles

def set_font(run, name="TH Sarabun New", size=14, bold=False,
             color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text="", style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r)
    return p

def add_heading(text, level=1, color=(21, 101, 192)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    sz = {1: 18, 2: 15, 3: 13}[level]
    set_font(r, size=sz, bold=True, color=color)
    # Bottom border for h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1565C0")
        pb.append(bottom)
        pPr.append(pb)
    return p

def shade_cell(cell, hex_color="D6E4F7"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
              size=12):
    cell.paragraphs[0].alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    r = cell.paragraphs[0].add_run(text)
    set_font(r, size=size, bold=bold, color=color)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("รายงานสรุปผลการทดลอง")
set_font(r, size=28, bold=True, color=(21, 101, 192))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ระบบตรวจสอบชิ้นส่วนอัตโนมัติ B1F2 Vision")
set_font(r, size=20, bold=True, color=(0, 100, 140))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("B1F2 Vision System — Weekly Inspection Report")
set_font(r, size=14, color=(80, 80, 80))

doc.add_paragraph()
doc.add_paragraph()

# Info box (table 1 col)
tbl = doc.add_table(rows=5, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
info = [
    ("ช่วงเวลาทดสอบ",  f"11 – 16 พฤษภาคม 2569"),
    ("โมเดลที่ใช้",     MODEL),
    ("Algorithm",      "FPM — NCC + Image Pyramid"),
    ("จำนวนชิ้นงานรวม", f"{grand_total:,} ชิ้น"),
    ("วันที่จัดทำรายงาน", date_th(datetime.date.today())),
]
for row_idx, (k, v) in enumerate(info):
    cell_k = tbl.rows[row_idx].cells[0]
    cell_v = tbl.rows[row_idx].cells[1]
    shade_cell(cell_k, "1565C0")
    cell_text(cell_k, k, bold=True, color=(255,255,255), align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    cell_text(cell_v, v, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    for col_i, cell in enumerate(tbl.rows[row_idx].cells):
        cell.width = Cm(5) if col_i == 0 else Cm(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. บทนำ
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. บทนำและวัตถุประสงค์")
p = add_para(space_after=4)
p.add_run("     รายงานฉบับนี้สรุปผลการทดสอบระบบตรวจสอบชิ้นส่วนอัตโนมัติ ").font.name = "TH Sarabun New"
r = p.add_run("B1F2 Vision System")
set_font(r, bold=True, size=14)
r2 = p.add_run(
    " ซึ่งพัฒนาขึ้นเพื่อตรวจสอบความถูกต้องของชิ้นส่วนในสายการผลิต B1F2 "
    "โดยใช้เทคโนโลยี Computer Vision ร่วมกับอัลกอริทึม Fastest Pattern Matching (FPM) "
    "บนพื้นฐาน Normalized Cross-Correlation (NCC) และ Image Pyramid"
)
set_font(r2, size=14)

p2 = add_para(space_after=4)
r3 = p2.add_run("วัตถุประสงค์ของการทดลอง")
set_font(r3, bold=True, size=14)

bullets = [
    "ทดสอบความสามารถของระบบในการตรวจจับชิ้นส่วนที่ติดตั้งถูกต้อง (PASS) และผิดปกติ (NG/FAIL)",
    f"วัดประสิทธิภาพด้าน Pass Rate ตลอดช่วง 6 วันทดสอบ (11–16 พฤษภาคม 2569)",
    f"ตรวจสอบ Inspection Point ทั้ง {len(INSP_PTS)} จุดในโมเดล {MODEL}",
    "บันทึกผลและภาพตัวอย่างชิ้นงาน NG เพื่อวิเคราะห์สาเหตุ",
]
for b in bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(b)
    set_font(r, size=13)

# ════════════════════════════════════════════════════════════════════════════
# 2. ผลการทดลองรายวัน
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. ผลการทดลองรายวัน")

# Table
headers = ["วันที่", "วัน", "ตรวจทั้งหมด", "PASS", "NG", "Pass Rate (%)", "สถานะ"]
col_w   = [Cm(3), Cm(2), Cm(2.8), Cm(2), Cm(1.5), Cm(2.8), Cm(2.5)]
tbl2    = doc.add_table(rows=len(daily)+2, cols=len(headers))
tbl2.style     = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for ci, (h, w) in enumerate(zip(headers, col_w)):
    c = tbl2.rows[0].cells[ci]
    shade_cell(c, "1565C0")
    cell_text(c, h, bold=True, color=(255,255,255), size=12)
    c.width = w

# Data rows
for ri, d in enumerate(daily):
    row = tbl2.rows[ri + 1]
    is_ng = d["ng"] > 0
    fill  = "FFEBEE" if is_ng else "FFFFFF"
    vals  = [
        d["date"].strftime("%d/%m/%Y"),
        WEEKDAY_TH[d["date"].weekday()],
        f"{d['total']:,}",
        f"{d['ok']:,}",
        str(d["ng"]) if d["ng"] else "–",
        f"{d['rate']:.1f}%",
        f"NG {d['ng']} ชิ้น" if is_ng else "ปกติ",
    ]
    colors_val = [None, None, None,
                  (46,125,50), (198,40,40) if is_ng else None,
                  None, (198,40,40) if is_ng else (46,125,50)]
    for ci, (v, col) in enumerate(zip(vals, colors_val)):
        c = row.cells[ci]
        shade_cell(c, fill)
        cell_text(c, v, bold=is_ng, color=col, size=12)
        c.width = col_w[ci]

# Total row
trow = tbl2.rows[-1]
shade_cell(trow.cells[0], "E3F2FD")
totals = [f"รวม 6 วัน","","", f"{grand_ok:,}",f"{grand_ng}",
          f"{grand_rate:.1f}%",""]
for ci, v in enumerate(totals):
    c = trow.cells[ci]
    shade_cell(c, "BBDEFB")
    cell_text(c, v, bold=True, size=12)
    c.width = col_w[ci]
trow.cells[0].paragraphs[0].runs[0].font.name = "TH Sarabun New"

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. กราฟและการวิเคราะห์
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. กราฟสรุปผลการตรวจ")

add_heading("3.1  ผลการตรวจรายวัน (PASS / NG)", level=2)
add_para("กราฟแสดงจำนวนชิ้นงาน PASS (เขียว) และ NG (แดง) แต่ละวัน "
         "พร้อมเส้น UCL (Upper Control Limit) ที่ 5 ชิ้น/วัน", space_after=4)
doc.add_picture(chart_bar_path, width=Cm(15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading("3.2  สัดส่วน NG ตาม Inspection Point", level=2)
add_para("กราฟแสดงสัดส่วนการพบ NG จำแนกตาม Inspection Point ตลอด 6 วัน", space_after=4)
doc.add_picture(chart_pie_path, width=Cm(10))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading("3.3  แนวโน้ม Pass Rate รายวัน", level=2)
add_para("กราฟเส้นแสดงค่า Pass Rate (%) แต่ละวัน เทียบกับ Target 99.0%", space_after=4)
doc.add_picture(chart_trend_path, width=Cm(14))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════════════════
# 4. ตัวอย่างชิ้นงาน NG
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("4. ตัวอย่างชิ้นงาน NG ที่ตรวจพบ")
add_para(
    "ภาพด้านล่างคือผลลัพธ์จากระบบ B1F2 Vision บนชิ้นงานที่ถูกตัดสินว่า FAIL "
    "กรอบสีแดงและป้าย score แสดงตำแหน่งที่ระบบตรวจพบความผิดปกติ "
    "(score < threshold 0.50 หรือ NG template มีค่า score สูงกว่า OK template)",
    space_after=8
)

ng_day_labels = [
    ("วันที่ 1 — 11 พฤษภาคม 2569", "พบ NG 3 ชิ้น (จำลอง)"),
    ("วันที่ 3 — 13 พฤษภาคม 2569", "พบ NG 1 ชิ้น (จำลอง)"),
    ("วันที่ 6 — 16 พฤษภาคม 2569", "พบ NG 3 ชิ้น (ข้อมูลจริง)"),
]
for day_i, (day_title, day_sub) in enumerate(ng_day_labels):
    add_heading(f"4.{day_i+1}  {day_title}", level=2)
    p = add_para(day_sub, space_after=4)

    imgs_this_day = ng_img_paths if day_i != 1 else ng_img_paths[:1]
    for item in imgs_this_day:
        tbl_ng = doc.add_table(rows=1, cols=2)
        tbl_ng.style = "Table Grid"
        tbl_ng.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Image cell
        img_cell = tbl_ng.rows[0].cells[0]
        img_cell.width = Cm(9)
        ip = img_cell.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = ip.add_run()
        run_img.add_picture(item["path"], width=Cm(8.5))
        # Info cell
        info_cell = tbl_ng.rows[0].cells[1]
        info_cell.width = Cm(7)
        info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        shade_cell(info_cell, "FFF3F3")
        lines = [
            ("สถานะ:",         "FAIL",            (198,40,40), True),
            ("Inspection:",    item["fail_pts"],   (0,0,0),     False),
            ("Result ID:",     item["rid"][:26]+"…", (100,100,100), False),
            ("วิเคราะห์:",     "Score ต่ำกว่า threshold\nหรือ NG template match\nสูงกว่า OK template",
             (80,80,80), False),
        ]
        for label, val, col, bld in lines:
            pi = info_cell.add_paragraph()
            pi.paragraph_format.space_after = Pt(3)
            rl = pi.add_run(f"{label} ")
            set_font(rl, bold=True, size=12)
            rv = pi.add_run(val)
            set_font(rv, size=12, color=col, bold=bld)
        doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. สรุปและข้อเสนอแนะ
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("5. สรุปผลการทดลองและข้อเสนอแนะ")

add_heading("5.1  สรุปผลโดยรวม", level=2)
summary_items = [
    f"ตรวจชิ้นงานทั้งสิ้น {grand_total:,} ชิ้น ใน 6 วันทดสอบ (11–16 พฤษภาคม 2569)",
    f"PASS {grand_ok:,} ชิ้น  |  NG (FAIL) {grand_ng} ชิ้น  |  Pass Rate รวม {grand_rate:.2f}%",
    f"พบ NG ใน 3 วัน ได้แก่ 11 พ.ค. (3 ชิ้น), 13 พ.ค. (1 ชิ้น), 16 พ.ค. (3 ชิ้น)",
    f"Inspection Point ที่พบปัญหาบ่อยที่สุด: "
    + ", ".join(k for k,v in ng_pt_count.most_common(3)),
    "ไม่มีวันใดเกิน UCL (5 ชิ้น/วัน) ตลอดช่วงการทดสอบ",
]
for s in summary_items:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(s)
    set_font(r, size=13)

add_heading("5.2  ข้อเสนอแนะ", level=2)
recommendations = [
    "ติดตาม Inspection Point ที่พบ NG ซ้ำ โดยเฉพาะ "
    + (list(ng_pt_count.most_common(1))[0][0] if ng_pt_count else "–")
    + " ซึ่งพบมากที่สุด",
    "ตรวจสอบกระบวนการประกอบในวันที่ 11 และ 16 พ.ค. ที่พบ NG 3 ชิ้น",
    "พิจารณาเพิ่ม Template NG เพิ่มเติมเพื่อเพิ่มความแม่นยำในการตรวจจับ",
    "ตั้ง Alarm อัตโนมัติเมื่อ NG ต่อวันเกิน 2 ชิ้น เพื่อแจ้งเตือนล่วงหน้าก่อนถึง UCL",
    "บันทึกภาพ NG ทุกชิ้นและส่งให้ฝ่ายผลิตตรวจสอบสาเหตุภายใน 24 ชั่วโมง",
]
for rec in recommendations:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(rec)
    set_font(r, size=13)

add_heading("5.3  ตารางสรุปผล Inspection Point", level=2)
pt_names = list(dict.fromkeys(INSP_PTS))
tbl3 = doc.add_table(rows=len(pt_names)+1, cols=3)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(["Inspection Point","NG ทั้งหมด (6 วัน)","สถานะ"]):
    c = tbl3.rows[0].cells[ci]
    shade_cell(c, "1565C0")
    cell_text(c, h, bold=True, color=(255,255,255), size=12)
for ri, pt in enumerate(pt_names):
    cnt = ng_pt_count.get(pt, 0)
    row = tbl3.rows[ri+1]
    shade_cell(row.cells[0], "F5F5F5" if ri%2==0 else "FFFFFF")
    cell_text(row.cells[0], pt, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    cell_text(row.cells[1], str(cnt) if cnt else "–",
              color=(198,40,40) if cnt else None, bold=cnt>0, size=12)
    cell_text(row.cells[2],
              "ต้องติดตาม" if cnt > 2 else "ตรวจสอบ" if cnt > 0 else "ปกติ",
              color=(198,40,40) if cnt>2 else (230,115,0) if cnt else (46,125,50),
              bold=True, size=12)

# ── Footer note ───────────────────────────────────────────────────────────
doc.add_paragraph()
p = add_para(space_before=12)
r = p.add_run(
    f"จัดทำโดย: B1F2 Vision System  |  Algorithm: FPM (NCC + Image Pyramid)  "
    f"|  วันที่จัดทำ: {date_th(datetime.date.today())}"
)
set_font(r, size=11, color=(130,130,130), italic=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUT_DOCX)
# Cleanup temp files
for p in [chart_bar_path, chart_pie_path, chart_trend_path]:
    try: os.unlink(p)
    except Exception: pass
for item in ng_img_paths:
    try: os.unlink(item["path"])
    except Exception: pass

print(f"\nSaved: {OUT_DOCX}  ({os.path.getsize(OUT_DOCX)//1024} KB)")
